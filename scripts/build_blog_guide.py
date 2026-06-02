from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "YT-Bookmarker-Blog-Guide.docx"
ICON = ROOT / "assets" / "icon-128.png"

RED = "E62117"
DEEP_RED = "B51D14"
INK = "172033"
MUTED = "5D6678"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FDEBE9"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_width.set(qn("w:type"), "dxa")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_table(table, widths_dxa: list[int]) -> None:
    set_table_geometry(table, widths_dxa)
    for index, row in enumerate(table.rows):
        for cell in row.cells:
            if index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
                    run.font.size = Pt(9.5)
                    if index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(INK)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_paragraph(doc, text="", *, style=None, after=6, before=0, color=INK, bold=False, italic=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_font(run, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_font(run, color=INK)


def add_number(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_font(run, color=INK)


def add_callout(doc, label: str, text: str, fill=LIGHT_RED) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, color=DEEP_RED, bold=True)
    text_run = paragraph.add_run(text)
    set_font(text_run, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F3F5")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code)
    set_font(run, size=9, color="263238", name="Courier New")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("YT BOOKMARKER  |  PRODUCT AND ENGINEERING GUIDE")
    set_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("YT Bookmarker  |  Page ")
    set_font(run, size=8.5, color=MUTED)
    add_page_number(footer)


def add_cover(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(20)
    paragraph.add_run().add_picture(str(ICON), width=Inches(1.25))

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("PRODUCT BUILD GUIDE")
    set_font(run, size=10, color=DEEP_RED, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Building YT Bookmarker")
    set_font(run, size=29, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("A professional Chrome extension for timestamp bookmarks and automatic YouTube replay-based Hot Moments")
    set_font(run, size=14, color=MUTED, italic=True)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(16)
    run = rule.add_run("BOOKMARKS  |  HEAT-MAP EXTRACTION  |  MANIFEST V3")
    set_font(run, size=9.5, color=DEEP_RED, bold=True)

    add_callout(
        doc,
        "Purpose",
        "Use this document as a technical reference, a product story, and a source outline for a blog post about building a polished browser extension.",
        LIGHT_RED,
    )

    add_paragraph(doc, "Inside this guide", before=10, after=5, color=INK, bold=True)
    for item in (
        "The product idea and user experience",
        "How YouTube heat-map extraction works",
        "Extension architecture and local storage",
        "Professional UI and icon decisions",
        "Testing strategy and Chrome Web Store readiness",
        "A suggested blog outline with reusable talking points",
    ):
        add_bullet(doc, item)

    doc.add_page_break()


def build_document() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. Product Story")
    add_paragraph(
        doc,
        "YT Bookmarker started with a simple problem: useful moments inside long YouTube videos are easy to forget. A timestamp bookmark solves the manual case, but YouTube already exposes another valuable signal: the replay heat map shown above the seek bar on many videos. The extension combines both ideas into one lightweight workflow.",
    )
    add_callout(
        doc,
        "Core promise",
        "Save the moments you choose and automatically capture the moments viewers replay most often.",
    )
    add_heading(doc, "What the user experiences", level=2)
    for item in (
        "A clean bookmark action appears alongside the native YouTube player controls.",
        "Clicking the bookmark action saves the current timestamp with video metadata.",
        "When YouTube exposes replay data, Hot Moments are extracted automatically after the player finishes loading.",
        "The popup clearly separates user-created bookmarks from automatic Hot Moments.",
        "Every saved moment can be played, edited, or deleted from the popup.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2. Extension Architecture")
    add_paragraph(
        doc,
        "The extension follows Chrome Manifest V3 conventions. Responsibilities are split so DOM work stays inside the YouTube page, persistence stays inside the service worker, and presentation stays inside the popup.",
    )
    architecture = doc.add_table(rows=1, cols=3)
    architecture.rows[0].cells[0].text = "Surface"
    architecture.rows[0].cells[1].text = "Primary file"
    architecture.rows[0].cells[2].text = "Responsibility"
    for surface, path, responsibility in (
        ("Content script", "src/content/index.ts", "Injects the bookmark control, observes YouTube player state, reads heat-map SVG data, and sends messages."),
        ("Service worker", "src/background/index.ts", "Creates bookmark records and coordinates local storage CRUD operations."),
        ("Popup", "src/popup/index.ts", "Loads the active video's data and renders Saved Bookmarks and Hot Moments as separate sections."),
        ("Analyzer", "src/core/analyzer.ts", "Parses SVG paths, calculates engagement scores, maps chapters, and clusters peaks."),
        ("Storage", "src/core/storage.ts", "Persists bookmarks under video-specific keys and prevents duplicate Hot Moments."),
    ):
        cells = architecture.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = surface, path, responsibility
    style_table(architecture, [1700, 2500, 5160])

    add_heading(doc, "Message flow", level=2)
    for step in (
        "The content script reads the current YouTube video ID and player metadata.",
        "A manual bookmark sends ADD_BOOKMARK. Automatic replay extraction sends ADD_HOT_MOMENTS.",
        "The service worker normalizes bookmark records and writes them through the storage module.",
        "The popup requests GET_BOOKMARKS_FOR_VIDEO for the active watch page.",
        "Popup controls send PLAY, UPDATE_BOOKMARK, or DELETE_BOOKMARK messages as needed.",
    ):
        add_number(doc, step)

    add_heading(doc, "3. Extracting YouTube Hot Moments")
    add_paragraph(
        doc,
        "The replay heat map is already rendered by YouTube as an SVG graph. The extension does not call a private API. It reads the graph that is present in the page and converts its shape into useful timestamps.",
    )
    add_heading(doc, "The DOM target", level=2)
    add_code(
        doc,
        '<div class="ytp-heat-map-container">\n'
        '  <div class="ytp-heat-map-chapter">\n'
        '    <svg class="ytp-heat-map-svg" viewBox="0 0 1000 100">\n'
        '      <path class="ytp-modern-heat-map" d="M 0.0,100.0 C ..." />\n'
        "    </svg>\n"
        "  </div>\n"
        "</div>",
    )
    add_paragraph(
        doc,
        "The important element is path.ytp-modern-heat-map. Its d attribute contains SVG path commands. Horizontal positions map to playback time; vertical positions represent replay intensity.",
    )

    add_heading(doc, "The extraction algorithm", level=2)
    for step in (
        "Find .ytp-heat-map-container and its .ytp-heat-map-chapter children.",
        "Read path.ytp-modern-heat-map and parse M, L, C, H, and V path commands.",
        "Use each command endpoint as a sample point from the rendered curve.",
        "Convert the vertical coordinate into a normalized replay score: (height - y) / height.",
        "Convert the horizontal coordinate into a timestamp using video duration and chapter position.",
        "Keep local maxima above the sensitivity threshold.",
        "Cluster timestamps that are too close together and keep the strongest candidate.",
        "Store up to five Hot Moments, ordered by replay score.",
    ):
        add_number(doc, step)

    add_heading(doc, "Why delayed loading matters", level=2)
    add_paragraph(
        doc,
        "YouTube is a single-page application. The video element, duration, controls, and heat-map SVG may arrive at different times. The extension watches player metadata events and heat-map DOM mutations, then retries extraction for a bounded period. This avoids racing the page during initial load or in-app navigation.",
    )
    add_callout(
        doc,
        "Important reliability fix",
        "The extractor does not require the seek bar to be visible. A hidden control layer can report a rendered width of zero even when the SVG path is valid. Single-chapter videos map directly across the full duration; multi-chapter videos fall back to YouTube's inline left and width styles.",
        LIGHT_BLUE,
    )

    add_heading(doc, "4. Bookmark Data And Duplicate Protection")
    add_paragraph(
        doc,
        "Bookmarks are stored locally with chrome.storage.local. Each record includes the video ID, timestamp, title, channel, description, creation time, source type, and optional replay score.",
    )
    add_code(
        doc,
        'type Bookmark = {\n'
        "  id: string;\n"
        "  videoId: string;\n"
        "  time: number;\n"
        "  title: string;\n"
        "  channel: string;\n"
        "  desc: string;\n"
        '  source?: "manual" | "hot-moment";\n'
        "  score?: number;\n"
        "};",
    )
    add_paragraph(
        doc,
        "Hot Moments are written in a batch. Before storage, the extension checks whether another Hot Moment already exists within five seconds. That makes automatic extraction idempotent when a page reloads or a mutation fires more than once.",
    )

    add_heading(doc, "5. Professional UI Decisions")
    add_heading(doc, "Player integration", level=2)
    add_paragraph(
        doc,
        "The player keeps one focused action: save the current timestamp. Hot Moments are automatic, so a second manual extraction action would add clutter and create uncertainty about whether the feature is working.",
    )
    add_heading(doc, "Popup sections", level=2)
    add_paragraph(
        doc,
        "The popup deliberately separates user intent from automated discovery:",
    )
    for item in (
        "Saved Bookmarks contains timestamps created by the user.",
        "Hot Moments contains replay peaks extracted automatically from YouTube.",
        "Each section has its own count, explanation, and empty state.",
        "Hot Moment cards include an engagement-score badge for context.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Extension icon", level=2)
    icon_paragraph = doc.add_paragraph()
    icon_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    icon_paragraph.paragraph_format.space_after = Pt(8)
    icon_paragraph.add_run().add_picture(str(ICON), width=Inches(0.85))
    add_paragraph(
        doc,
        "The icon combines a bookmark shape, a replay curve, and a small highlight spark. A post-processing script removes only the connected outer white canvas, crops the artwork tightly, and exports transparent PNG variants at 16, 32, 48, and 128 pixels. This prevents a visible white frame in the Chrome toolbar.",
    )
    add_code(doc, "python scripts/process_extension_icon.py SOURCE_IMAGE assets")

    add_heading(doc, "6. Testing Strategy")
    add_paragraph(
        doc,
        "The automated suite focuses on the parsing and mapping logic because that is the feature with the highest risk of silent failure. The production bundle is also checked with TypeScript and esbuild.",
    )
    tests = doc.add_table(rows=1, cols=2)
    tests.rows[0].cells[0].text = "Test area"
    tests.rows[0].cells[1].text = "Coverage"
    for area, coverage in (
        ("SVG parsing", "Cubic endpoints, relative coordinates, and scientific notation."),
        ("YouTube fixture", "A production-shaped M ... C ... replay path like the one rendered in YouTube."),
        ("Ranking", "The strongest replay moments appear first."),
        ("Chapter mapping", "Chapter-local SVG coordinates map to full-video timestamps."),
        ("Clustering", "Nearby peaks collapse into the strongest timestamp."),
        ("Hidden controls", "Single-chapter extraction works even when rendered width is zero."),
        ("Inline geometry", "Multi-chapter extraction falls back to YouTube left and width styles."),
        ("Invalid input", "Invalid durations produce no Hot Moments."),
    ):
        cells = tests.add_row().cells
        cells[0].text, cells[1].text = area, coverage
    style_table(tests, [2200, 7160])
    add_heading(doc, "Verification commands", level=2)
    add_code(doc, "yarn test\nyarn typecheck\nyarn build")

    add_heading(doc, "7. Local Development Workflow")
    for step in (
        "Install dependencies with yarn install.",
        "Run yarn test, yarn typecheck, and yarn build.",
        "Open chrome://extensions and enable Developer mode.",
        "Choose Load unpacked and select the repository folder.",
        "After rebuilding, reload the extension card and refresh an open YouTube watch page.",
        "Test on a video with a replay heat map, then verify the two popup sections.",
    ):
        add_number(doc, step)

    add_heading(doc, "8. Chrome Web Store Readiness")
    add_paragraph(
        doc,
        "The extension is intentionally small and local-first. Before publishing, complete the store packaging work below.",
    )
    for item in (
        "Create final store screenshots that show the player bookmark action and the split popup sections.",
        "Write a concise privacy policy explaining that bookmarks are stored locally in chrome.storage.local.",
        "Review the requested permissions: storage, tabs, and YouTube host access.",
        "Prepare a store description focused on timestamp recall and automatic replay highlights.",
        "Test navigation across several YouTube layouts, chapters, long videos, and videos without heat maps.",
        "Package a clean release build and confirm version metadata in manifest.json.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "9. Suggested Blog Outline")
    add_paragraph(
        doc,
        "The following outline can be used as the skeleton for a technical blog post.",
    )
    outline = (
        ("Hook", "Long videos contain valuable moments, but returning to them is harder than it should be."),
        ("The product idea", "Combine intentional timestamp bookmarks with YouTube's replay heat map."),
        ("The DOM discovery", "Show the ytp-heat-map-container hierarchy and explain the SVG path."),
        ("The algorithm", "Walk through parsing, score calculation, timestamp mapping, and clustering."),
        ("The browser-extension architecture", "Explain content script, service worker, popup, and storage boundaries."),
        ("The reliability lesson", "Discuss delayed SPA rendering and the hidden-controls width-zero bug."),
        ("The polish pass", "Explain why one player action, split popup sections, and a transparent toolbar icon matter."),
        ("Testing", "Share the production-shaped fixture and the regression cases."),
        ("What comes next", "Mention search, tags, export, sync, and store submission."),
    )
    for heading, detail in outline:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.25
        title_run = paragraph.add_run(f"{heading}: ")
        set_font(title_run, color=DEEP_RED, bold=True)
        detail_run = paragraph.add_run(detail)
        set_font(detail_run, color=INK)

    add_heading(doc, "10. Future Enhancements")
    for item in (
        "Search across all bookmarked videos.",
        "Tags, collections, and study-oriented workflows.",
        "Import and export for backup or sharing.",
        "Keyboard shortcuts for fast timestamp capture.",
        "Optional cloud sync with a clear privacy model.",
        "Video thumbnails and richer metadata in an all-videos dashboard.",
        "More browser-compatibility testing before expanding beyond Chrome.",
    ):
        add_bullet(doc, item)

    add_callout(
        doc,
        "Closing thought",
        "The strongest part of YT Bookmarker is not the amount of code. It is the way a small extension turns an existing visual signal into a useful, automatic workflow while keeping the manual bookmark experience simple.",
        LIGHT_BLUE,
    )

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
